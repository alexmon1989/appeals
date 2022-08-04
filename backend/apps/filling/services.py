from django.contrib.auth import get_user_model
from django.utils.datastructures import MultiValueDict
from django.http.request import QueryDict
from django.db.models import Prefetch, Count, Q
from django.db.models.query import QuerySet
from django.core.files.base import ContentFile
from django.conf import settings

from elasticsearch import Elasticsearch
from elasticsearch_dsl import Search, Q as Q_Es

from docx import Document as PyDocxDocument
from docx.shared import Pt
from docxcompose.composer import Composer

from ..classifiers.models import DocumentType
from ..cases.models import Document, Sign, DocumentTemplate
from ..cases.services import services as cases_services
from .models import ClaimField, Claim
from ..common.utils import docx_replace
from ..common.utils import base64_to_temp_file

from typing import List, Type, Union
from pathlib import Path
from distutils.dir_util import copy_tree
import json
import datetime
import tempfile
import uuid
import os


UserModel = get_user_model()


def claim_process_input_data(data: dict) -> dict:
    """Обрабатывает данные заявки, пришедшие от пользователя."""
    for key in data:
        if 'date' in key:
            try:
                data[key] = datetime.datetime.strptime(data[key], '%d.%m.%Y').strftime('%Y-%m-%d')
            except:
                pass
    return data


def claim_create(post_data: QueryDict, files_data: dict, user: UserModel) -> Claim:
    """Создаёт обращение пользователя."""
    stage_3_field = ClaimField.objects.filter(claim_kind=post_data['claim_kind'], stage=3).first().input_id
    claim = Claim.objects.create(
        obj_kind_id=post_data['obj_kind'],
        claim_kind_id=post_data['claim_kind'],
        third_person=post_data.get('third_person', False),
        obj_number=post_data[stage_3_field],
        obj_title=post_data['obj_title'],
        json_data=json.dumps(claim_process_input_data(post_data)),
        user=user
    )

    # Загрузка файлов
    stage_9_fields = ClaimField.objects.filter(
        claim_kind=post_data['claim_kind'],
        stage=9,
        field_type__in=(ClaimField.FieldType.FILE, ClaimField.FieldType.FILE_MULTIPLE)
    )
    for field in stage_9_fields:
        if field.input_id in files_data or f"{field.input_id}[]" in files_data:
            # Получение типа документа
            doc_type, created = DocumentType.objects.get_or_create(title=field.title)
            # Сохранение файла
            if field.field_type == ClaimField.FieldType.FILE:
                # Создание документа
                doc = Document.objects.create(
                    claim=claim,
                    document_type=doc_type,
                    input_date=datetime.datetime.now(),
                    claim_document=True,
                )

                file_data = files_data[field.input_id][0]

                # Сохранение файла во временный каталог
                tmp_file_path = base64_to_temp_file(file_data['content'])

                # Сохранение файла в постоянный каталог и в БД
                document_save_file(doc, file_data['name'], tmp_file_path, True)

                if doc_type.base_doc:
                    # Создание файла обращения (с шапкой)
                    document_create_main_claim_doc_file(claim, doc)
            else:
                files = files_data[f"{field.input_id}[]"]

                for file_data in files:
                    # Создание документа
                    doc = Document.objects.create(
                        claim=claim,
                        document_type=doc_type,
                        input_date=datetime.datetime.now(),
                        claim_document=True
                    )

                    # Сохранение файла во временный каталог
                    tmp_file_path = base64_to_temp_file(file_data['content'])

                    # Сохранение файла в постоянный каталог и в БД
                    document_save_file(doc, file_data['name'], tmp_file_path, True)

    return claim


def claim_edit(сlaim_id: int, post_data: dict, files_data: MultiValueDict, user: UserModel) -> Union[Claim, None]:
    """Редактирует обращение пользователя."""
    claim = claim_get_user_claims_qs(user).filter(pk=сlaim_id, status__lt=3).first()
    if not claim:
        return None

    # Удаление документов, которые указал пользователь и тех, которые генерируются автоматически
    try:
        docs_to_delete_ids = post_data['delete_doc_id']
        if type(docs_to_delete_ids) is not list:
            docs_to_delete_ids = [docs_to_delete_ids]
        del post_data['delete_doc_id']
    except KeyError:
        Document.objects.filter(claim=claim, auto_generated=True).delete()
    else:
        Document.objects.filter(Q(pk__in=docs_to_delete_ids) | Q(claim=claim, auto_generated=True)).delete()

    stage_3_field = ClaimField.objects.filter(claim_kind=post_data['claim_kind'], stage=3).first().input_id

    # Обновление текстовых данных обращения
    claim.obj_kind_id = post_data['obj_kind']
    claim.claim_kind_id = post_data['claim_kind']
    claim.third_person = post_data.get('third_person', False)
    claim.obj_number = post_data[stage_3_field]
    claim.obj_title = post_data['obj_title']
    claim.json_data = json.dumps(claim_process_input_data(post_data))
    claim.user = user
    claim.save()

    # Загрузка файлов
    stage_9_fields = ClaimField.objects.filter(
        claim_kind=post_data['claim_kind'],
        stage=9,
        field_type__in=(ClaimField.FieldType.FILE, ClaimField.FieldType.FILE_MULTIPLE)
    )
    for field in stage_9_fields:
        if field.input_id in files_data or f"{field.input_id}[]" in files_data:
            # Получение типа документа
            doc_type, created = DocumentType.objects.get_or_create(title=field.title)

            # Сохранение файла
            if field.field_type == ClaimField.FieldType.FILE:
                # Удаление "старого" документа этого типа
                Document.objects.filter(claim=claim, document_type=doc_type).delete()

                # Создание документа
                doc = Document.objects.create(
                    claim=claim,
                    document_type=doc_type,
                    input_date=datetime.datetime.now(),
                    claim_document=True,
                )

                file_data = files_data[field.input_id][0]

                # Сохранение файла во временный каталог
                tmp_file_path = base64_to_temp_file(file_data['content'])

                # Сохранение файла в постоянный каталог и в БД
                document_save_file(doc, file_data['name'], tmp_file_path, True)
            else:
                files = files_data[f"{field.input_id}[]"]

                for file_data in files:
                    # Создание документа
                    doc = Document.objects.create(
                        claim=claim,
                        document_type=doc_type,
                        input_date=datetime.datetime.now(),
                        claim_document=True
                    )

                    # Сохранение файла во временный каталог
                    tmp_file_path = base64_to_temp_file(file_data['content'])

                    # Сохранение файла в постоянный каталог и в БД
                    document_save_file(doc, file_data['name'], tmp_file_path, True)

    # Формирование основного документа обращения (заявления)
    base_doc = Document.objects.filter(claim=claim, document_type__base_doc=True).first()
    if base_doc:
        document_create_main_claim_doc_file(claim, base_doc)
        claim.status = 1
        claim.save()

    return claim


def claim_get_fields(bool_as_int: bool = False) -> List[ClaimField]:
    """Вовзращает возможные поля обращений (зависящие от типа)"""
    claim_fields = list(ClaimField.objects.order_by('pk').filter(enabled=True).values(
        'pk',
        'title',
        'help_text',
        'input_id',
        'field_type',
        'allowed_extensions',
        'stage',
        'claim_kind_id',
        'editable',
        'required',
    ))
    if bool_as_int:
        for claim_field in claim_fields:
            claim_field['editable'] = int(claim_field['editable'])
            claim_field['required'] = int(claim_field['required'])

    return claim_fields


def claim_get_user_claims_qs(user: UserModel) -> QuerySet[Claim]:
    """Возвращает обращения пользователя."""
    return Claim.objects.filter(
        user=user
    ).select_related(
        'claim_kind', 'obj_kind', 'case'
    ).prefetch_related(
        'document_set', 'document_set__document_type'
    )


def claim_get_stages_details(claim: Claim) -> dict:
    """Возвращает данные заявки по этапам ввода формы."""
    fields = ClaimField.objects.filter(claim_kind=claim.claim_kind, stage__lt=9)
    stages = {
        3: {
            'title': 'Номер заявки/охоронного документа',
            'items': [],
        },
        4: {
            'title': 'Дані об\'єкта права інтелектуальної власності',
            'items': [],
        },
        5: {
            'title': 'Відомості про заявника або власника',
            'items': [],
        },
        6: {
            'title': 'Відомості про апелянта',
            'items': [],
        },
        7: {
            'title': 'Додаткова інформація',
            'items': [],
        },
        8: {
            'title': 'Відомості щодо рішення Укрпатенту',
            'items': [],
        },
    }

    claim_data = json.loads(claim.json_data)
    for field in fields:
        try:
            if claim_data[field.input_id]:
                stages[field.stage]['items'].append({
                    'title': field.title,
                    'value': claim_data[field.input_id],
                    'type': field.field_type,
                    'id': field.input_id,
                })
        except KeyError:
            pass

    if stages[3]['items'][0]['id'] == 'app_number':
        stages[5]['title'] = 'Відомості про заявника'
    else:
        stages[5]['title'] = 'Відомості про власника'

    return stages


def claim_get_documents_qs(claim_id: int) -> QuerySet[Document]:
    """Возвращает список документов обращения."""
    documents = Document.objects.filter(
        claim_id=claim_id,
        claim_document=True
    ).select_related(
        'document_type'
    ).prefetch_related(
        Prefetch('sign_set', queryset=Sign.objects.all())
    ).order_by(
        '-auto_generated',
        'document_type'
    ).annotate(Count('sign'))

    return documents


def claim_get_documents(claim_id: int) -> list:
    """Возвращает список документов обращения в формате json."""
    documents = claim_get_documents_qs(claim_id)
    res = []
    for document in documents:
        res.append({
            'id': document.pk,
            'document_type': document.document_type.title,
            'auto_generated': int(document.auto_generated),
            'file_url': document.file.url,
            'file_name': Path(document.file.name).name,
            'sign__count': document.sign__count,
        })
    return res


def claim_set_status_if_all_docs_signed(claim_id: Type[int]) -> None:
    """Меняет статус обращения с "1" на "2" (готово к рассмотрению), если все документы подписаны пользователем."""
    try:
        claim = Claim.objects.get(pk=claim_id, status=1)
    except Claim.DoesNotExist:
        pass
    else:
        # Проверка все ли документы, поданные пользователем подписаны.
        documents = Document.objects.filter(
            claim=claim,
            claim_document=True
        ).prefetch_related(
            Prefetch(
                'sign_set',
                queryset=Sign.objects.filter(user_id=claim.user_id)
            )
        )
        for document in documents:
            # Какой-то документ не подписан
            if document.sign_set.count() == 0:
                return

        # Обновление статуса
        claim.status = 2
        claim.save()


def claim_get_data_by_id(claim_id: int, user: UserModel = None, **kwargs) -> dict:
    """Возвращает данные обращения пользователя."""
    if user:
        claim = claim_get_user_claims_qs(user).filter(pk=claim_id, **kwargs).first()
    else:
        claim = Claim.objects.filter(pk=claim_id, **kwargs).first()
    if claim:
        res = {
            'claim_data': {
                'id': claim.pk,
                'obj_number': claim.obj_number,
                'obj_kind': claim.obj_kind.title,
                'obj_kind_id': claim.obj_kind.id,
                'claim_kind': claim.claim_kind.title,
                'claim_kind_id': claim.claim_kind.id,
                'status_verbal': claim.get_status_display(),
                'status': claim.status,
                'submission_date': claim.submission_date.strftime('%d.%m.%Y %H:%M:%S'),
            },
            'stages': claim_get_stages_details(claim),
            'documents': claim_get_documents(claim)
        }

        if claim.status == 3:
            res['claim_data']['case_number'] = claim.case.case_number

        return res
    return {}


def claim_copy_docs_to_external_server(claim_id: int) -> None:
    """Копирует документы с внутреннего сервера на внешний """
    documents = Document.objects.filter(claim_id=claim_id)
    for doc in documents:
        # Каталог с документом на внутреннем сервере
        doc_folder_relative = Path(str(doc.file)).parent
        from_folder_path = settings.MEDIA_ROOT / doc_folder_relative

        # Каталог с документом на внешнем сервере
        to_folder_path = settings.EXTERNAL_MEDIA_ROOT / doc_folder_relative

        copy_tree(str(from_folder_path), str(to_folder_path))


def claim_create_files_with_signs_info(claim_id: int, signs: list) -> None:
    documents = Document.objects.filter(claim_id=claim_id)
    for doc in documents:
        if doc.sign_set.count() == 0:
            cases_services.document_add_sign_info_to_file(doc.pk, signs)


def document_get_data_for_main_claim_doc_file(claim_id: Type[int]) -> dict:
    """Возвращает словарь с данными, необходимыми для формирования основного документа обращения."""
    claim = Claim.objects.filter(pk=claim_id).first()
    claim_data = json.loads(claim.json_data)

    res = {
        '{{ OBJ_KIND_TITLE }}': claim.obj_kind.title,
        '{{ CLAIM_KIND_TITLE }}': claim.claim_kind.template_title,
        '{{ OBJ_TITLE }}': claim_data['obj_title'],
        '{{ APP_NUMBER }}': claim_data['app_number'],
        '{{ APP_DATE }}': claim_data['app_date'],
        '{{ DECISION_DATE }}': claim_data.get('decision_date', ''),
        '{{ DECISION_RECEIVE_DATE }}': claim_data.get('decision_receive_date', ''),
        '{{ APPLICANT_TITLE }}': claim_data.get('applicant_title', '').replace('\r\n', '\n'),
        '{{ REGISTRATION_NUMBER }}': claim_data.get('registration_number'),
        '{{ OWNER_TITLE }}': claim_data.get('owner_title', '').replace('\r\n', '\n'),
        '{{ OWNER_ADDRESS }}': claim_data.get('owner_address', '').replace('\r\n', '\n'),
        '{{ WELL_KNOWN_DATE }}': claim_data.get('well_known_date'),
        '{{ WELL_KNOWN_CLASSES }}': claim_data.get('well_known_classes', '').replace('\r\n', '\n'),
    }

    for key in ('{{ APP_DATE }}', '{{ DECISION_DATE }}', '{{ DECISION_RECEIVE_DATE }}', '{{ WELL_KNOWN_DATE }}'):
        if res[key]:
            res[key] = datetime.datetime.strptime(res[key], '%Y-%m-%d').strftime('%d.%m.%Y')

    if claim.third_person:
        res['{{ APPELAINT_TITLE }}'] = claim_data.get('third_person_applicant_title', '').replace('\r\n', '\n')
        res['{{ APPELAINT_ADDRESS }}'] = claim_data.get('third_person_applicant_address', '').replace('\r\n', '\n')
        res['{{ REPRESENT_TITLE }}'] = claim_data.get('third_person_represent_title', '').replace('\r\n', '\n')
        res['{{ REPRESENT_ADDRESS }}'] = claim_data.get('third_person_represent_address', '').replace('\r\n', '\n')
        res['{{ MAILING_ADDRESS }}'] = claim_data.get('third_person_mailing_address', '').replace('\r\n', '\n')
        res['{{ CONTACTS_PHONE }}'] = claim_data.get('third_person_phone', '')
        res['{{ CONTACTS_EMAIL }}'] = claim_data.get('third_person_email', '')
    else:
        res['{{ APPELAINT_TITLE }}'] = claim_data.get('applicant_title', '').replace('\r\n', '\n') \
                                       or claim_data.get('owner_title', '').replace('\r\n', '\n')
        res['{{ APPELAINT_ADDRESS }}'] = claim_data.get('applicant_address', '').replace('\r\n', '\n') \
                                         or claim_data.get('owner_address', '').replace('\r\n', '\n')
        res['{{ REPRESENT_TITLE }}'] = claim_data.get('represent_title', '').replace('\r\n', '\n')
        res['{{ REPRESENT_ADDRESS }}'] = claim_data.get('represent_address', '').replace('\r\n', '\n')
        res['{{ MAILING_ADDRESS }}'] = claim_data.get('mailing_address', '').replace('\r\n', '\n')
        res['{{ CONTACTS_PHONE }}'] = claim_data.get('phone', '')
        res['{{ CONTACTS_EMAIL }}'] = claim_data.get('email', '')

    return res


def document_create_main_claim_doc_file(claim: Claim, base_doc: Document) -> Document:
    """Создаёт документ файл обращения."""
    doc_type = DocumentType.objects.filter(claim_kinds__id=claim.claim_kind.pk, create_with_claim=True).first()
    if doc_type:
        doc_template = DocumentTemplate.objects.filter(documents_types__code=doc_type.code).first()
        f_header = open(doc_template.file.path, 'rb')
        doc_header = PyDocxDocument(f_header)

        doc_data_to_replace = document_get_data_for_main_claim_doc_file(claim.pk)
        docx_replace(doc_header, doc_data_to_replace)

        tmp_dir = tempfile.gettempdir()
        tmp_file_name = f'{uuid.uuid4()}.docx'
        tmp_file_path = Path(tmp_dir) / tmp_file_name

        doc_header.save(tmp_file_path)

        master = doc_header
        composer = Composer(master)

        f = open(base_doc.file.path, 'rb')
        doc_body = PyDocxDocument(f)
        # doc_body = PyDocxDocument(base_doc.file)
        composer.append(doc_body)
        composer.save(tmp_file_path)
        f.close()
        f_header.close()

        # Поставить всему документу 12-й размер шрифта и Times New Roman
        input_doc = PyDocxDocument(tmp_file_path)
        for paragraph in input_doc.paragraphs:
            paragraph.style = input_doc.styles['Normal']
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        input_doc.save(tmp_file_path)

        doc = Document.objects.create(
            claim=claim,
            document_type=doc_type,
            input_date=datetime.datetime.now(),
            claim_document=True,
            auto_generated=1
        )

        with open(tmp_file_path, "rb") as fh:
            with ContentFile(fh.read()) as file_content:
                doc.file.save(tmp_file_name, file_content)
                doc.save()

        os.remove(tmp_file_path)

        return doc


def application_is_published(data: dict) -> bool:
    """Определяет, опубликована ли заявка."""
    if data['search_data']['obj_state'] == 2:  # охранные документы всегда опубликованы
        return True

    if data['Document']['idObjType'] == 1:  # изобретения
        return 'I_43.D' in data['Claim']

    if data['Document']['idObjType'] == 4:  # торговые марки
        date_441 = data['TradeMark']['TrademarkDetails'].get('Code_441')
        if date_441:
            return True
        mark_status = application_fixed_mark_status_code(data)
        app_date = datetime.datetime.strptime(data['search_data']['app_date'][:10], '%Y-%m-%d')
        date_441_start = datetime.datetime.strptime('2020-08-18', '%Y-%m-%d')
        return mark_status > 2000 and app_date < date_441_start

    if data['Document']['idObjType'] in (2, 3, 5, 6):  # остальные типы объектов не публикуются
        return False


def application_user_belongs_to_app(data: dict, user_names: list) -> bool:
    """Определяет, принадлежит пользователь к заявке."""
    allowed_persons = []

    for person_type in ('applicant', 'inventor', 'owner', 'agent'):
        try:
            allowed_persons += [x['name'] for x in data['search_data'][person_type]]
        except KeyError:
            pass

    # Адреса для листування (ТМ)
    try:
        allowed_persons.append(
            data['TradeMark']['TrademarkDetails']['CorrespondenceAddress']['CorrespondenceAddressBook'][
                'Name']['FreeFormatNameLine']
        )
    except KeyError:
        pass

    # Адреса для листування (ПЗ)
    try:
        allowed_persons.append(
            data['Design']['DesignDetails']['CorrespondenceAddress']['CorrespondenceAddressBook'][
                'FormattedNameAddress']['Name']['FreeFormatName']['FreeFormatNameDetails']['FreeFormatNameLine']
        )
    except KeyError:
        pass

    # Адреса для листування (заявки на винаходи, корисні моделі)
    try:
        allowed_persons.append(data['Claim']['I_98'])
    except KeyError:
        pass

    # Адреса для листування (охоронні документи на винаходи, корисні моделі)
    try:
        allowed_persons.append(data['Patent']['I_98'])
    except KeyError:
        pass

    # Удаление None-объектов
    allowed_persons = list(filter(lambda item: item is not None, allowed_persons))

    # Замена латинских I, i на кириллицу
    user_names = [x.replace('I', 'І').replace('i', 'і') for x in user_names]
    allowed_persons = [x.replace('I', 'І').replace('i', 'і') for x in allowed_persons]

    # Проверка на вхождение
    return any([person for person in allowed_persons for user_name in user_names if user_name.upper() in person.upper()])


def application_fixed_mark_status_code(data: dict) -> int:
    """Анализирует список документов и возвращает код статуса согласно их наличию."""
    result = int(data['Document'].get('MarkCurrentStatusCodeType', 0))
    for doc in data['TradeMark'].get('DocFlow', {}).get('Documents', []):
        if ('ТM-1.1' in doc['DocRecord']['DocType'] or 'ТМ-1.1' in doc['DocRecord']['DocType']) and result < 2000:
            result = 3000
        if ('Т-05' in doc['DocRecord']['DocType'] or 'Т-5' in doc['DocRecord']['DocType']) and result < 3000:
            result = 3000
        if 'Т-08' in doc['DocRecord']['DocType'] and result < 4000:
            result = 4000
    return result


def application_get_data_from_es(obj_num_type: str, obj_number: str, obj_kind_id_sis: int, obj_state: int) -> dict:
    """Получает данные заявки из ElasticSearch СИС."""
    client = Elasticsearch(settings.ELASTIC_HOST, timeout=settings.ELASTIC_TIMEOUT)

    if obj_num_type == 'registration_number':
        obj_num_type = 'protective_doc_number'

    query_string = f"search_data.{obj_num_type}.exact:{obj_number} AND " \
                   f"Document.idObjType:{obj_kind_id_sis} AND " \
                   f"search_data.obj_state:{obj_state}"

    q = Q_Es(
        'query_string',
        query=query_string
    )

    s = Search(using=client, index=settings.ELASTIC_INDEX_NAME).query(q).execute()

    if not s:
        return {}
    return s[0].to_dict()


def document_save_file(doc: Document, file_name: str, tmp_file_path: Path, remove_file: bool = False) -> None:
    """Сохраняет файл документа с диска."""
    # Сохранение файла в постоянный каталог
    with open(tmp_file_path, "rb") as fh:
        with ContentFile(fh.read()) as file_content:
            doc.file.save(file_name, file_content)
            doc.save()

    # Удаление временного файла
    if remove_file:
        os.remove(tmp_file_path)
