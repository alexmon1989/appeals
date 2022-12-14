import datetime

from django.contrib.auth import get_user_model
from django.conf import settings
from django.utils import timezone

from apps.cases.models import Document, Sign, DocumentHistory, Command, PostalProtocolExchange, EsignProtocolExchange
from apps.cases.utils import set_cell_border
from apps.common.utils import (docx_replace, generate_barcode_img, substitute_image_docx)
from apps.classifiers.models import DocumentType, CommandType

from typing import List
from pathlib import Path
from docx import Document as DocumentWord
import random
from typing import Union
import subprocess

UserModel = get_user_model()


def document_get_by_id(doc_id: int) -> Document:
    """Возвращает документ по его идентификатору."""
    doc = Document.objects.filter(pk=doc_id).prefetch_related(
        'sign_set',
        'sign_set__user',
        'documenthistory_set',
        'documenthistory_set__user'
    ).select_related(
        'document_type',
        'case',
        'claim',
    )
    return doc.first()


def document_add_sign_info_to_file(doc_id: int,
                                   signs: list,
                                   internal_document: bool = False,
                                   user_id: int = None) -> None:
    """Создаёт .docx с инф-ей о номере документа, штрихкоде
    и цифровых подписях в конце (на основе оригинального файла документа)"""
    document = document_get_by_id(doc_id)

    if not document.registration_number:
        # Документ создался во внутреннем модуле
        if internal_document:
            # Присвоение документу номера, штрихкода, даты регистрации в момент подписания
            document_set_reg_number(doc_id)
            document_set_barcode(doc_id)
            document_set_reg_date(doc_id)
            document.refresh_from_db()

        if Path(str(document.file)).suffix == '.docx':
            # Открытие документа
            docx_file_path = Path(settings.MEDIA_ROOT) / Path(str(document.file))
            docx_with_signs_file_path = docx_file_path.parent / f"{docx_file_path.stem}_signs.docx"
            docx = DocumentWord(docx_file_path)

            # Документ создался во внутреннем модуле
            if internal_document:
                # Добавление номера документа, даты регистрации и штрих-кода в файл
                file_vars = {
                    '{{ DOC_REG_DATE }}': document.registration_date.strftime('%d.%m.%Y'),
                    '{{ DOC_REG_NUM }}': document.registration_number,
                }
                docx_replace(docx, file_vars)
                barcode_file_path = generate_barcode_img(document.barcode)
                substitute_image_docx(docx, '{{ BARCODE_IMG }}', barcode_file_path, 6)
                document_add_history(
                    document.pk,
                    'Документу присвоєні номер та штрих-код (автоматично в момент підпису).',
                    user_id
                )

            # Добавление таблички с информацией о подписях (только если один подписант)
            if document.sign_set.count() == 1:
                docx.add_paragraph('Підписали:')
                table = docx.add_table(rows=len(signs), cols=1)
                for i, sign in enumerate(signs):
                    cells = table.rows[i].cells
                    cells[0].paragraphs[0].add_run(
                        f"{sign['subject']}\n{sign['serial_number']}\n{sign['issuer']}"
                    )
                    paragraph_format = cells[0].paragraphs[0].paragraph_format
                    paragraph_format.space_after = 0
                    # set_cell_margins(cells[0], top=50, start=50, bottom=50, end=50)
                    set_cell_border(
                        cells[0],
                        top={"sz": 12, "val": "single", "color": "black", "space": "0"},
                        bottom={"sz": 12, "val": "single", "color": "black", "space": "0"},
                        start={"sz": 12, "val": "single", "color": "black", "space": "0"},
                        end={"sz": 12, "val": "single", "color": "black", "space": "0"},
                    )
                document_add_history(
                    document.pk,
                    'Створено новий документ з інформацією про підписантів (автоматично)',
                    user_id
                )

            # Сохранение
            docx.save(docx_with_signs_file_path)

            if internal_document:
                # Конвертация в pdf
                subprocess.call(
                    ["libreoffice", "--headless", '--convert-to', 'pdf', docx_with_signs_file_path],
                    cwd=str(Path(docx_with_signs_file_path).parent)
                )
                document.converted_to_pdf = True
                document.save()
                document_add_history(document.pk, 'Конвертовано у pdf (автоматично)', user_id)


def document_set_reg_number(doc_id: int) -> None:
    """Присваивает документу регистрационный номер, который возвращает глобальный нумератор."""
    numbers = ''.join([str(random.randint(0, 9)) for _ in range(5)])
    document = Document.objects.get(pk=doc_id)
    if document.case:
        document.registration_number = f"Вих-{numbers}/2022"
    else:
        document.registration_number = f"Вх-{numbers}/2022"
    document.save()


def document_set_reg_date(doc_id: int) -> None:
    """Присваивает документу дату регистрации."""
    document = Document.objects.get(pk=doc_id)
    document.registration_date = datetime.datetime.now()
    document.save()


def document_set_barcode(doc_id: int) -> None:
    """Присваивает документу штрихкод, который возвращает глобальный нумератор."""
    document = Document.objects.get(pk=doc_id)
    document.barcode = ''.join([str(random.randint(0, 9)) for _ in range(32)])
    document.save()


def document_get_signs_info(doc_id: int) -> List[dict]:
    """Возвращает информацию о подписавших документ."""
    res = []
    signs = Sign.objects.filter(
        document_id=doc_id
    ).order_by(
        'created_at'
    ).values(
        'subject',
        'serial_number',
        'issuer',
        'timestamp'
    )
    for sign in signs:
        res.append(sign)
    return res


def document_can_be_signed_by_user(document: Union[int, Document], user: UserModel) -> bool:
    """Проверяет, может ли пользователь подписывать документ."""
    if isinstance(document, int):
        document = document_get_by_id(document)
    if document.case:
        for sign in document.sign_set.all():
            if sign.user_id == user.pk:
                # Документ подписан пользователем
                if sign.timestamp:
                    return False
                # Документ ожидает подписания
                return True
        return False

    # Заявка (могут подписывать только заявители, которые подавали)
    return document.claim and document.claim.user_id == user.pk and user.is_applicant


def document_get_case_documents_to_sign(case_id: int, user: UserModel) -> list:
    """Возвращает список документов дела, которые ожидают подписания пользователем."""
    documents = Document.objects.filter(
        sign__user=user,
        sign__timestamp='',
        case_id=case_id
    ).select_related('document_type').order_by('-created_at')

    res = []

    for document in documents:
        res.append({
            'id': document.pk,
            'document_type': document.document_type.title,
            'auto_generated': int(document.auto_generated),
            'file_url': document.file.url,
            'file_name': Path(document.file.name).name,
        })

    return res


def document_add_history(doc_id: int, action: str, user_id: int = None) -> None:
    """Создаёт запись в историю дела."""
    DocumentHistory.objects.create(
        document_id=doc_id,
        action=action,
        user_id=user_id
    )


def document_soft_delete(doc_id: int, user: UserModel) -> Union[Document, None]:
    """Удаляет документ (мягкое удаление)."""
    document = document_get_by_id(doc_id)
    if document and document.can_be_deleted(user):
        document.deleted = True
        document.save()
        document_add_history(doc_id, 'Видалено', user.pk)
        return document
    return None


def document_convert_original_doc_to_pdf(document: Document, user_id: int) -> None:
    """Конвертирует оригинальный файл документа в PDF (также ставит метку в БД, что документ конвертирован)."""
    docx_file_path = Path(settings.MEDIA_ROOT) / Path(str(document.file))
    subprocess.call(
        ["libreoffice", "--headless", '--convert-to', 'pdf', docx_file_path],
        cwd=str(docx_file_path.parent)
    )
    # Переименование pdf
    rename_from = docx_file_path.parent / f"{docx_file_path.stem}.pdf"
    rename_to = docx_file_path.parent / f"{docx_file_path.stem}_signs.pdf"
    rename_from.rename(rename_to)

    document.converted_to_pdf = True
    document.save()
    document_add_history(document.pk, 'Конвертовано у pdf (автоматично)', user_id)


def document_send_to_sign_collegium(document: Document, user_id: int) -> None:
    """Передаёт документ на подпись членам коллегии и секретарю."""
    for item in document.case.collegiummembership_set.all():
        # Члены коллегии
        Sign.objects.create(
            document=document,
            user=item.person,
        )
    # Секретарь дела
    Sign.objects.get_or_create(
        document=document,
        user=document.case.secretary,
    )
    document_add_history(
        document.pk,
        'Документ відправлено на підпис членам апеляційної колегії та секретарю',
        user_id
    )


def document_send_to_sign_collegium_head(document: Document, user_id: int) -> None:
    """Передаёт документ на подпись главе коллегии."""
    Sign.objects.create(
        document=document,
        user=document.case.collegium_head,
    )
    document_add_history(
        document.pk,
        'Документ відправлено на підпис голові апеляційної колегії',
        user_id
    )


def document_send_to_director(document: Document, user_id: int) -> None:
    """Передаёт документ на подпись главе коллегии."""
    document_send_to_external_sign(document.pk)
    Sign.objects.get_or_create(
        document=document,
        file_signed=document.signed_file,
        external_service_sign=True
    )
    document_add_history(document.pk, 'Документ відправлено на підпис директору організації', user_id)


def document_send_to_sign(doc_id: int, user_id: int) -> None:
    """Создаёт записи для подписи пользователей."""
    document = document_get_by_id(doc_id)

    # Конвертация документа в pdf
    if not document.converted_to_pdf:
        document_convert_original_doc_to_pdf(document, user_id)

    # Передача на подпись в зависимости от типа подписанта
    sign_methods = {
        DocumentType.SignerType.COLLEGIUM.value: document_send_to_sign_collegium,
        DocumentType.SignerType.COLLEGIUM_HEAD.value: document_send_to_sign_collegium_head,
        DocumentType.SignerType.DIRECTOR.value: document_send_to_director,
    }
    sign_methods[document.document_type.signer_type](document, user_id)


def document_send_to_chancellary(pk: int, user_id: int) -> Command:
    """Отправляет документ в АС Вихідні документи, фактически - создаёт записи в БД."""
    command_type = CommandType.objects.get(command_name='send_to_chancellary')
    now = timezone.now()

    command = Command()
    command.command_type = command_type
    command.document_id = pk
    command.create_date = now
    command.is_done = False
    command.save()

    protocol = PostalProtocolExchange()
    protocol.command = command
    protocol.doc_id = pk
    protocol.id_cead = 0
    protocol.saved_at = now
    protocol.save()

    document_add_history(pk, 'Документ відправлено в АС "Вихідні документи"', user_id)

    return command


def document_send_to_external_sign(pk: int) -> Command:
    """Передаёт документ на подпись во внешний сервис.
    Фактически делает записи в таблицах БД для работы стороннего приложения,
    реально отправляющего документ на подпись."""
    command_type = CommandType.objects.get(command_name='send_to_esign_service')
    now = timezone.now()

    command = Command()
    command.command_type = command_type
    command.document_id = pk
    command.create_date = now
    command.is_done = False
    command.save()

    protocol = EsignProtocolExchange()
    protocol.command = command
    protocol.doc_id = pk
    protocol.id_cead = 0
    protocol.saved_at = now
    protocol.save()

    document_add_history(pk, 'Документ відправлено до сервісу підписання документів.')

    return command
