from django.contrib.auth import get_user_model

from apps.cases.models import Document, Case
from apps.classifiers.models import DocumentType

from .document_services import document_add_history
from apps.common.utils import docx_replace, first_lower, get_random_file_name, get_temp_file_path

from docx import Document as PyDocxDocument

from pathlib import Path

UserModel = get_user_model()


class Service:
    case: Case
    document: Document
    doc_type: DocumentType
    user_id: int
    signer: UserModel
    file_vars = {}
    extra_args = {}

    def _get_file_vars(self) -> dict:
        """Формирует список с переменными для замены в файле docx."""
        if self.doc_type.code == '0005':
            return get_file_vars_0005(self.case, self.signer)
        elif self.doc_type.code == '0006':
            return get_file_vars_0006(self.case, self.document)
        elif self.doc_type.code == '0007':
            return get_file_vars_0007(self.case, self.document)
        elif self.doc_type.code == '0009':
            return get_file_vars_0009(self.case, self.document)
        elif self.doc_type.code == '0010':
            return get_file_vars_0010(self.case, self.document)
        elif self.doc_type.code == '0011':
            return get_file_vars_0011(self.case, self.document)
        elif self.doc_type.code in ('0012', '0013', '0014', '0015', '0016', '0017', '0018', '0019', '0020', '0021',
                                    '0022', '0023'):
            return get_file_vars_stopping(self.case, self.extra_args['form_data'])
        elif self.doc_type.code in ('0024', '0025', '0026'):
            return get_file_vars_meeting(self.case, self.document)
        elif self.doc_type.code == '0027':
            return get_file_vars_pre_meeting_protocol(self.case, self.document)
        else:
            return {}

    def _create_doc_file(self) -> Path:
        """Создаёт файл на диске."""
        # Открытие файла с шаблоном
        docx = PyDocxDocument(self.doc_type.template.path)

        # Получение и замена переменных в файле
        docx_replace(docx, self.file_vars)

        # def iter_target_paragraphs(document):
        #     """Generate each paragraph inside all tables of `document`."""
        #     for table in document.tables:
        #         for row in table.rows:
        #             for cell in row.cells:
        #                 for paragraph in cell.paragraphs:
        #                     yield paragraph
        #
        # def substitute_image_placeholder(paragraph, image_var, barcode_file):
        #     # --- start with removing the placeholder text ---
        #     paragraph.text = paragraph.text.replace(image_var, "")
        #     # --- then append a run containing the image ---
        #     run = paragraph.add_run()
        #     from docx.shared import Cm
        #     run.add_picture(barcode_file, height=Cm(6))
        #
        # for paragraph in iter_target_paragraphs(docx):
        #     if '{{ BARCODE_IMG }}' in paragraph.text:
        #         substitute_image_placeholder(paragraph, '{{ BARCODE_IMG }}', str(barcode_file_path))

        # Сохранение файла во временный каталог
        tmp_file_name = get_random_file_name('docx')
        tmp_file_path = get_temp_file_path(tmp_file_name)
        docx.save(tmp_file_path)
        return tmp_file_path

    def _create_doc_db(self) -> Document:
        """Создаёт документ в БД (без ссылки на файл)."""
        doc = Document.objects.create(
            case=self.case,
            document_type=self.doc_type,
            # input_date=datetime.datetime.now(),
            auto_generated=True,
        )
        document_add_history(
            doc.pk,
            'Документ додано у систему (створено автоматично)',
            self.user_id
        )
        # document_set_barcode(doc.pk)
        # document_set_reg_number(doc.pk)
        doc.refresh_from_db()
        return doc

    def execute(self, case_id: int, doc_code: str, user_id: int, signer_id: int = None, **kwargs) -> Document:
        self.user_id = user_id
        self.signer = UserModel.objects.get(pk=signer_id) if signer_id else None
        self.extra_args = kwargs
        self.case = Case.objects.select_related('claim').get(pk=case_id)
        self.doc_type = DocumentType.objects.get(
            code=doc_code,
            claim_kinds__id=self.case.claim.claim_kind_id
        )
        self.document = self._create_doc_db()
        self.file_vars = self._get_file_vars()
        doc_file_path = self._create_doc_file()
        self.document.assign_file(doc_file_path)
        return self.document


def get_file_vars_0005(case: Case, signer: UserModel) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0005
    (Розпорядження про створення колегії)."""
    # Члены коллегии
    common_members = []
    head_title = ''
    head_position = ''
    for member in case.collegiummembership_set.all():
        if member.is_head:
            head_title = member.person.get_full_name
            head_position = first_lower(member.person.position)
        else:
            common_members.append(
                {
                    'title': member.person.get_full_name,
                    'position': first_lower(member.person.position)
                }
            )

    # Апелянт
    appellant = case.claim.get_appellant_title()

    case_extra_info = f"апелянт - {appellant}"
    if case.claim.claim_kind.claim_sense == 'DE':
        case_extra_info += f", номер заявки: {case.claim.obj_number}"

    if case.claim.third_person:
        if case.claim.claim_kind.claim_sense == 'DE':
            case_extra_info += f", заявник: {case.claim.get_applicant_title()}"
        else:
            case_extra_info += f"номер охоронного документа: {case.claim.obj_number}, " \
                               f"власник: {case.claim.get_owner_title()}"

    return {
        '{{ CASE_NUMBER }}': case.case_number,
        '{{ CLAIM_KIND }}': first_lower(case.claim.claim_kind.template_title),
        '{{ OBJ_KIND }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ CASE_EXTRA_INFO }}': case_extra_info,
        '{{ HEAD_TITLE }}': head_title,
        '{{ HEAD_POSITION }} ': head_position,
        '{{ MEMBER_1_TITLE }}': common_members[0]['title'],
        '{{ MEMBER_1_POSITION }}': common_members[0]['position'],
        '{{ MEMBER_2_TITLE }}': common_members[1]['title'],
        '{{ MEMBER_2_POSITION }}': common_members[1]['position'],
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_POSITION }}': case.secretary.position,
        '{{ SIGNER_TITLE }}': signer.get_full_name,
        '{{ SIGNER_POSITION }}': signer.position,
    }


def get_file_vars_0006(case: Case, document: Document) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0006
    (сообщение апеллянту о принятии дела к рассмотрению)."""
    # Представитель апеллянта или апеллянт
    represent = case.claim.get_represent_title(third_person=case.claim.third_person)
    if represent:
        header_person_title = represent
        header_person_address = case.claim.get_represent_address()
    else:
        header_person_title = case.claim.get_appellant_title()
        header_person_address = case.claim.get_appellant_address()

    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    # Коллегия
    collegium_head = ''
    collegium = []
    for item in case.collegiummembership_set.select_related('person').all():
        collegium.append(item.person.get_full_name_initials())
        if item.is_head:
            collegium_head = item.person.get_full_name
    return {
        '{{ HEADER_PERSON_TITLE }}': header_person_title,
        '{{ HEADER_PERSON_ADDRESS }}': header_person_address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ APP_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ APPLICANT_TITLE }}': case.claim.get_applicant_title(),
        '{{ APPLICANT_ADDRESS }}': case.claim.get_applicant_address(),
        '{{ COLLEGIUM_MEMBERS }}': ', '.join(collegium),
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ COLLEGIUM_HEAD }}': collegium_head,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',
    }


def get_file_vars_0007(case: Case, document: Document) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0007
    (сообщение заявителю о принятии дела к рассмотрению)."""

    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    # Коллегия
    collegium = []
    for item in case.collegiummembership_set.all():
        collegium.append(item.person.get_full_name_initials())
    collegium_head = case.collegium_head.get_full_name

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ APP_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ APPLICANT_TITLE }}': case.claim.get_applicant_title(),
        '{{ APPLICANT_ADDRESS }}': case.claim.get_applicant_address(),
        '{{ COLLEGIUM_MEMBERS }}': ', '.join(collegium),
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ COLLEGIUM_HEAD }}': collegium_head,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',

        '{{ DOC_DOWNLOAD_CODE }}': claim_doc.barcode[-10:],  # последние 10 цифр штрих-кода
    }


def get_file_vars_0009(case: Case, document: Document) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0009 -
    Повідомлення апелянту про прийняття апеляційної заяви до розгляду"""
    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    # Коллегия
    collegium = []
    for item in case.collegiummembership_set.all():
        collegium.append(item.person.get_full_name_initials())
    collegium_head = case.collegium_head.get_full_name

    # Подготовительное заседание
    pre_meeting = case.meeting_set.filter(meeting_type='PRE').first()

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ REGISTRATION_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ COLLEGIUM_MEMBERS }}': ', '.join(collegium),
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ COLLEGIUM_HEAD }}': collegium_head,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',
        '{{ PRE_MEETING_DATETIME }}': pre_meeting.datetime.strftime('%d.%m.%Y %H:%M:%S') if pre_meeting else '',

        '{{ DOC_DOWNLOAD_CODE }}': claim_doc.barcode[-10:],  # последние 10 цифр штрих-кода
    }


def get_file_vars_0010(case: Case, document: Document) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0010 -
    Повідомлення власнику про прийняття апеляційної заяви до розгляду"""
    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    # Коллегия
    collegium = []
    for item in case.collegiummembership_set.all():
        collegium.append(item.person.get_full_name_initials())
    collegium_head = case.collegium_head.get_full_name

    # Подготовительное заседание
    pre_meeting = case.meeting_set.filter(meeting_type='PRE').first()

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ REGISTRATION_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ OWNER_TITLE }}': case.claim.get_owner_title(),
        '{{ OWNER_ADDRESS }}': case.claim.get_owner_address(),
        '{{ COLLEGIUM_MEMBERS }}': ', '.join(collegium),
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ COLLEGIUM_HEAD }}': collegium_head,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',
        '{{ PRE_MEETING_DATETIME }}': pre_meeting.datetime.strftime('%d.%m.%Y %H:%M:%S') if pre_meeting else '',

        '{{ DOC_DOWNLOAD_CODE }}': claim_doc.barcode[-10:],  # последние 10 цифр штрих-кода
    }


def get_file_vars_0011(case: Case, document: Document) -> dict:
    """Возвращает словарь со значениями переменных для формирования документа с кодом 0011 -
    Повідомлення апелянту про прийняття заяви про визнання ТМ ДВ до розгляду"""

    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    # Коллегия
    collegium = []
    for item in case.collegiummembership_set.all():
        collegium.append(item.person.get_full_name_initials())
    collegium_head = case.collegium_head.get_full_name

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ APP_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ COLLEGIUM_MEMBERS }}': ', '.join(collegium),
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ COLLEGIUM_HEAD }}': collegium_head,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',

        '{{ DOC_DOWNLOAD_CODE }}': claim_doc.barcode[-10:],  # последние 10 цифр штрих-кода
    }


def get_file_vars_stopping(case: Case, form_data: dict):
    """Переменные для формирования файла документа оповещения об остановке рассмотрения дела / признания непригодным."""

    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ APP_NUMBER }}': case.claim.obj_number,
        '{{ REGISTRATION_NUMBER }}': case.claim.obj_number,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ APPLICANT_TITLE }}': case.claim.get_applicant_title(),
        '{{ APPLICANT_ADDRESS }}': case.claim.get_applicant_address(),
        '{{ COLLEGIUM_HEAD }}': case.collegium_head.get_full_name,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',
        '{{ REASON }}': form_data.get('reason', ''),
        '{{ CIRCUMSTANCES }}': form_data.get('circumstances', ''),
    }


def get_file_vars_meeting(case: Case, document: Document):
    """Переменные для формирования файла документа оповещения об назначении заседания."""
    # Документ обращения
    claim_doc = Document.objects.get(document_type__code__in=('0001', '0002', '0003', '0004'), claim=case.claim)
    claim_doc_reg_num = claim_doc.registration_number
    claim_doc_reg_date = claim_doc.input_date.strftime("%d.%m.%Y")

    return {
        '{{ HEADER_PERSON_TITLE }}': case.addressee,
        '{{ HEADER_PERSON_ADDRESS }}': case.address,
        '{{ CLAIM_DOC_REG_NUM }}': claim_doc_reg_num,
        '{{ CLAIM_DOC_REG_DATE }}': claim_doc_reg_date,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ APPELAINT_TITLE }}': case.claim.get_appellant_title(),
        '{{ APPELAINT_ADDRESS }}': case.claim.get_appellant_address(),
        '{{ COLLEGIUM_HEAD }}': case.collegium_head.get_full_name,
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ SECRETARY_PHONE }}': case.secretary.phone_number or '',
        '{{ SECRETARY_EMAIL }}': case.secretary.email,
        '{{ MEETING_DATETIME }}': case.meeting_set.order_by('-pk').first().datetime.strftime('%d.%m.%Y %H:%M:%S'),
    }


def get_file_vars_pre_meeting_protocol(case: Case, document: Document):
    """Переменные для формирования файла документа оповещения об назначении заседания."""
    # Коллегия
    collegium = []
    for item in case.collegiummembership_set.filter(is_head=False):
        collegium.append(item.person)

    return {
        '{{ MEETING_DATE }}': case.meeting_set.order_by('-pk').first().datetime.strftime('%d.%m.%Y %H:%M:%S'),
        '{{ CASE_NUMBER }}': case.case_number,
        '{{ COLLEGIUM_HEAD }}': case.collegium_head.get_full_name,
        '{{ COLLEGIUM_MEMBERS }}': f'{collegium[0].get_full_name}, {collegium[1].get_full_name}',
        '{{ SECRETARY_TITLE }}': case.secretary.get_full_name,
        '{{ OBJ_NUMBER }}': case.claim.obj_number,
        '{{ OBJ_KIND_TITLE }}': first_lower(case.claim.obj_kind.title),
        '{{ OBJ_TITLE }}': case.claim.obj_title,
        '{{ COLLEGIUM_MEMBER_1 }}': collegium[0].get_full_name,
        '{{ COLLEGIUM_MEMBER_2 }}': collegium[1].get_full_name,
    }
