from django.contrib.auth import get_user_model
from django.conf import settings
from django.db.models import Count, Q, QuerySet

from ..models import Case, Document, Sign, CaseStage, CaseStageStep, CaseHistory
from ..utils import set_cell_border
from ...filling.models import Claim
from ...filling import services as filling_services

from typing import Iterable, List, Union
from pathlib import Path
from docx import Document as DocumentWord
import datetime
import random


UserModel = get_user_model()


def case_get_list() -> QuerySet[Case]:
    """Возвращает список апелляционных дел, к которым есть доступ у пользователя"""
    cases = Case.objects.select_related(
        'claim',
        'claim__obj_kind',
        'claim__claim_kind',
        'papers_owner',
        'expert',
        'secretary',
        'stage_step',
        'stage_step__stage',
    ).prefetch_related(
        'collegiummembership_set__person',
        'document_set',
        'document_set__document_type',
        'document_set__sign_set',
        'refusal_reasons'
    )

    return cases


def case_filter_dt_list(cases: QuerySet[Case], current_user_id: int, user: str = None, obj_kind: int = None,
                        stage: str = None) -> Iterable[Case]:
    """Фильтрует список ап. дел по определённым параметрам."""

    # Оставляет только ап. дела, к которым имеет отношение пользователь
    if user and user == 'me':
        current_user = UserModel.objects.get(pk=current_user_id)
        cases = cases.filter(case_get_user_has_access_filter(current_user))

    # Фильтр по типу ОИС
    if obj_kind and obj_kind != 'all':
        cases = cases.filter(claim__obj_kind__id=obj_kind)

    if stage and stage != 'all':
        if stage == 'new':
            cases = cases.filter(stage_step__code=1000)
        else:
            cases = cases.filter(stage_step__code__gt=1000)

    return cases


def case_user_has_access(case_id: int, user: UserModel = None) -> bool:
    """Проверяет, имеет ли доступ к апелляционному делу пользователь."""
    return user.is_privileged \
           or Case.objects.filter(
                Q(
                    pk=case_id
                )
                & case_get_user_has_access_filter(user)
            ).exists()


def case_get_user_has_access_filter(user: UserModel = None) -> Q:
    """Возвращает запрос для фильтра, определяющего имеет ли пользователь доступ к ап. делу(ам)."""
    return (
                Q(collegiummembership__person=user)
                | Q(expert=user)
                | Q(secretary=user)
                | Q(papers_owner=user)
        )


def case_get_documents_list(case_id: int) -> Iterable[Document]:
    """Возвращает документы дела."""
    queryset = Document.objects.filter(
        Q(case_id=case_id) | Q(claim_id=Case.objects.get(pk=case_id).claim_id)
    ).select_related(
        'document_type'
    ).prefetch_related(
        'sign_set'
    ).annotate(
        signs_count=Count('sign')
    ).order_by('-created_at')

    return queryset


def case_generate_next_number() -> str:
    """Генерирует следующий номер дела."""
    cur_year = datetime.datetime.now().year
    last_case = Case.objects.filter(created_at__year=cur_year).order_by('-created_at').first()
    if last_case:
        last_case_num = last_case.case_number.split('/')[0]
        last_case_num = str(int(last_case_num) + 1).zfill(4)
        return f"{last_case_num}/{cur_year}"
    else:
        return f"0001/{cur_year}"


def case_create_from_claim(claim_id: int, user: UserModel) -> Union[Case, None]:
    """Создаёт дело на основе обращения."""
    claim = filling_services.claim_get_user_claims_qs(user).filter(pk=claim_id, status=2).first()
    if claim:
        case = Case.objects.create(
            claim=claim,
            case_number=case_generate_next_number(),
            stage_step=CaseStageStep.objects.get(code=1000)
        )
        claim.status = 3
        claim.submission_date = datetime.datetime.now()
        claim.save()

        # Присваивание документам номеров и штрихкодов
        for doc in claim.document_set.all():
            document_set_reg_number(doc.pk)
            document_set_barcode(doc.pk)

        return case
    return None


def case_get_stages(case_id: int) -> Union[List[dict], None]:
    """Возвращает стадии ап. дела."""
    case = case_get_list().filter(pk=case_id).first()
    if case:
        stages = CaseStage.objects.order_by('number')
        res = []
        current_stage_step = case.stage_step
        for stage in stages:
            if current_stage_step.stage.number > stage.number:
                status = 'done'
            elif current_stage_step.stage.number == stage.number:
                if current_stage_step.case_stopped:
                    status = 'done'
                else:
                    status = 'current'
            else:
                status = 'not-active'

            res.append({
                'title': stage.title,
                'number': stage.number,
                'status': status
            })
        return res
    return None


def case_add_history_action(case_id: int, action: str, user_id: int) -> None:
    """Добавляет действие в историю действий дела."""
    CaseHistory.objects.create(
        case_id=case_id,
        action=action,
        user_id=user_id
    )


def case_take_to_work(case_id: int, user_id: int) -> bool:
    """Принимает дело в работу (назначает секретаря и меняет статус на 2000)."""
    case = Case.objects.filter(pk=case_id, stage_step__code=1000).first()
    if case:
        case.secretary_id = user_id
        case.save()
        case_change_stage_step(case_id, 2000, user_id)
        return True
    return False


def case_change_stage_step(case_id: int, stage_step_code: int, user_id: int) -> None:
    """Присваивает делу новую стадию и делает отметку в журнале дела."""
    case = Case.objects.get(pk=case_id)
    stage_step = CaseStageStep.objects.get(code=stage_step_code)
    case.stage_step = stage_step
    case.save()
    case_add_history_action(
        case_id,
        f'Зміна стадії справи на "{stage_step.title}" (код {stage_step.code})',
        user_id
    )


def document_get_by_id(doc_id: int) -> Document:
    """Возвращает документ по его идентификатору."""
    doc = Document.objects.filter(pk=doc_id).prefetch_related('sign_set', 'document_type')
    return doc.first()


def document_add_sign_info_to_file(doc_id: int, signs: list) -> None:
    """Создаёт .docx с инф-ей о цифровых подписях в конце (на основе оригинального файла документа)"""
    document = document_get_by_id(doc_id)

    if Path(str(document.file)).suffix == '.docx':
        # Открытие документа
        docx_file_path = Path(settings.MEDIA_ROOT) / Path(str(document.file))
        docx_with_signs_file_path = docx_file_path.parent / f"{docx_file_path.stem}_signs.docx"
        docx = DocumentWord(docx_file_path)

        # Добавление таблички с информацией о подписях
        docx.add_paragraph('Підписали:')
        table = docx.add_table(rows=len(signs), cols=1)
        for i, sign in enumerate(signs):
            cells = table.rows[i].cells
            cells[0].paragraphs[0].add_run(
                f"{sign['subject']}\n{sign['serial_number']}\n{sign['issuer']}"
            )
            paragraph_format = cells[0].paragraphs[0].paragraph_format
            paragraph_format.space_after = 0
            # set_cell_margins(cells[0], top=50, start=50, bottom=50, end=50)
            set_cell_border(
                cells[0],
                top={"sz": 12, "val": "single", "color": "black", "space": "0"},
                bottom={"sz": 12, "val": "single", "color": "black", "space": "0"},
                start={"sz": 12, "val": "single", "color": "black", "space": "0"},
                end={"sz": 12, "val": "single", "color": "black", "space": "0"},
            )

        # Сохранение
        docx.save(docx_with_signs_file_path)

    # Смена статуса обращения если все документы подписаны
    # filling_services.claim_set_status_if_all_docs_signed(document.claim_id)


def document_set_reg_number(doc_id: int) -> str:
    """Присваивает документу регистрационный номер, который возвращает глобальный нумератор."""
    if settings.DEBUG:
        numbers = ''.join([str(random.randint(0, 9)) for _ in range(5)])
        document = Document.objects.get(pk=doc_id)
        document.registration_number = f"Вх-{numbers}/2022"
        document.save()
    return ''


def document_set_barcode(doc_id: int) -> str:
    """Присваивает документу штрихкод, который возвращает глобальный нумератор."""
    if settings.DEBUG:
        document = Document.objects.get(pk=doc_id)
        document.barcode = ''.join([str(random.randint(0, 9)) for _ in range(32)])
        document.save()
    return ''


def document_get_signs_info(doc_id: int) -> List[dict]:
    """Возвращает информацию о подписавших документ."""
    res = []
    signs = Sign.objects.filter(
        document_id=doc_id
    ).order_by(
        'created_at'
    ).values(
        'subject',
        'serial_number',
        'issuer',
        'timestamp'
    )
    for sign in signs:
        res.append(sign)
    return res


def document_can_be_signed_by_user(document: Document, user: UserModel) -> bool:
    """Проверяет, может ли пользователь подписывать документ."""
    if Sign.objects.filter(document=document, user=user).exists():
        return False

    if document.claim:
        applicants = Claim.objects.filter(pk=document.claim.pk, user=user).exists()
    else:
        applicants = False

    if document.case:
        internal_users = document.document_type.title == 'Вихідний' and case_user_has_access(document.case_id, user)
    else:
        internal_users = False

    return applicants or internal_users


def sign_upload(file, dest: Path) -> bool:
    """Загружает файл с цифровой подписью на диск."""
    with open(dest, 'wb+') as destination:
        for chunk in file.chunks():
            destination.write(chunk)
    return True


def sign_create(data: dict) -> Sign:
    """Сохраняет данные о цифровой подписи в БД."""
    sign = Sign(**data)
    sign.save()
    return sign
